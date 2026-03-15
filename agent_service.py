import base64
import hashlib
import json
import os
import re
import shutil
import subprocess
import threading
import time
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from openai import APIConnectionError, APITimeoutError, APIError, RateLimitError

try:
    from docx2python import docx2python
except ImportError:
    docx2python = None

# XML namespaces needed for text-box extraction
_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_NS_V = "urn:schemas-microsoft-com:vml"

_GLM_BASE_URL = "https://open.bigmodel.cn/api/paas/v4/"
_LLM_TIMEOUT_SECONDS = 120
_LLM_MAX_RETRIES = 3
_LLM_MIN_REQUEST_INTERVAL_SECONDS = float(os.environ.get("GLM_MIN_REQUEST_INTERVAL_SECONDS", "3"))
_LLM_RATE_LIMIT_BACKOFF_SECONDS = float(os.environ.get("GLM_RATE_LIMIT_BACKOFF_SECONDS", "15"))
_ENABLE_VERIFICATION = os.environ.get("GLM_ENABLE_VERIFICATION", "false").lower() in {"1", "true", "yes", "on"}
_WORKFLOW_MODE = os.environ.get("GLM_WORKFLOW_MODE", "lite").lower()
_INCLUDE_HEADER_FOOTER = os.environ.get("GLM_INCLUDE_HEADER_FOOTER", "true").lower() in {"1", "true", "yes", "on"}
_ENABLE_GLM_CACHE = os.environ.get("GLM_ENABLE_CACHE", "true").lower() in {"1", "true", "yes", "on"}
_CACHE_DIR = Path(os.environ.get("GLM_CACHE_DIR", ".cache/glm"))
_llm_request_lock = threading.Lock()
_last_llm_request_at = 0.0


class GLMRateLimitError(RuntimeError):
    pass


class GLMConnectionError(RuntimeError):
    pass


def _build_llm(model_name: str, temperature: float = 0.2) -> ChatOpenAI:
    return ChatOpenAI(
        model=model_name,
        temperature=temperature,
        api_key=os.environ["GLM_API_KEY"],
        base_url=_GLM_BASE_URL,
        timeout=_LLM_TIMEOUT_SECONDS,
        max_retries=1,
    )


def _normalize_text(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").split("\n")).strip()


def _leading_indent(text: str) -> str:
    """Capture manual leading indentation characters used in templates."""
    match = re.match(r"^[\t \u3000]+", text or "")
    return match.group(0) if match else ""


def _preserve_leading_indent(original_text: str, rewritten_text: str) -> str:
    original_indent = _leading_indent(original_text)
    if not original_indent:
        return rewritten_text

    rewritten_indent = _leading_indent(rewritten_text)
    if rewritten_indent:
        return rewritten_text

    return f"{original_indent}{rewritten_text}"


def _serialize_for_cache(payload: Any) -> str:
    if isinstance(payload, str):
        return payload
    return json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)


def _get_model_name(llm: ChatOpenAI) -> str:
    return getattr(llm, "model_name", None) or getattr(llm, "model", "unknown-model")


def _get_cache_path(namespace: str, key_material: str) -> Path:
    digest = hashlib.sha256(key_material.encode("utf-8")).hexdigest()
    return _CACHE_DIR / namespace / f"{digest}.json"


def _read_cache(namespace: str, key_material: str) -> str | None:
    if not _ENABLE_GLM_CACHE:
        return None
    path = _get_cache_path(namespace, key_material)
    if not path.exists():
        return None
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None
    return str(payload.get("content", "")) if payload.get("content") else None


def _write_cache(namespace: str, key_material: str, content: str) -> None:
    if not _ENABLE_GLM_CACHE:
        return
    path = _get_cache_path(namespace, key_material)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps({"content": content}, ensure_ascii=False), encoding="utf-8")


def _extract_json(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        parts = cleaned.split("```")
        for part in parts:
            piece = part.strip()
            if piece.startswith("json"):
                piece = piece[4:].strip()
            if piece.startswith("{"):
                cleaned = piece
                break
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("Model did not return JSON content")
    return json.loads(cleaned[start : end + 1])


def _invoke_json(llm: ChatOpenAI, prompt: str) -> dict[str, Any]:
    content = _invoke_text(llm, prompt, cache_namespace="json")
    return _extract_json(content)


def _invoke_text(llm: ChatOpenAI, payload: Any, cache_namespace: str = "text") -> str:
    cache_key = f"{_get_model_name(llm)}\n{_serialize_for_cache(payload)}"
    cached = _read_cache(cache_namespace, cache_key)
    if cached is not None:
        return cached

    response = _invoke_llm_with_retry(llm, payload)
    content = response.content if isinstance(response.content, str) else str(response.content)
    _write_cache(cache_namespace, cache_key, content)
    return content


def _wait_for_llm_slot() -> None:
    global _last_llm_request_at
    with _llm_request_lock:
        now = time.monotonic()
        elapsed = now - _last_llm_request_at
        if elapsed < _LLM_MIN_REQUEST_INTERVAL_SECONDS:
            time.sleep(_LLM_MIN_REQUEST_INTERVAL_SECONDS - elapsed)
        _last_llm_request_at = time.monotonic()


def _is_rate_limit_error(exc: Exception) -> bool:
    message = str(exc).lower()
    return isinstance(exc, RateLimitError) or "429" in message or "1302" in message or "rate limit" in message or "请求频率" in message or "速率限制" in message


def _invoke_llm_with_retry(llm: ChatOpenAI, payload: Any):
    last_error: Exception | None = None
    for attempt in range(1, _LLM_MAX_RETRIES + 1):
        try:
            _wait_for_llm_slot()
            return llm.invoke(payload)
        except RateLimitError as exc:
            last_error = exc
            if attempt >= _LLM_MAX_RETRIES:
                break
            time.sleep(_LLM_RATE_LIMIT_BACKOFF_SECONDS * attempt)
        except (APIConnectionError, APITimeoutError, APIError) as exc:
            last_error = exc
            if _is_rate_limit_error(exc):
                if attempt >= _LLM_MAX_RETRIES:
                    break
                time.sleep(_LLM_RATE_LIMIT_BACKOFF_SECONDS * attempt)
                continue
            if attempt >= _LLM_MAX_RETRIES:
                break
            time.sleep(min(2 ** (attempt - 1), 4))
        except Exception as exc:
            message = str(exc).lower()
            if _is_rate_limit_error(exc):
                last_error = exc
                if attempt >= _LLM_MAX_RETRIES:
                    break
                time.sleep(_LLM_RATE_LIMIT_BACKOFF_SECONDS * attempt)
                continue
            if "connection error" not in message and "timed out" not in message:
                raise
            last_error = exc
            if attempt >= _LLM_MAX_RETRIES:
                break
            time.sleep(min(2 ** (attempt - 1), 4))

    detail = str(last_error) if last_error else "Unknown connection failure"
    if last_error and _is_rate_limit_error(last_error):
        raise GLMRateLimitError(
            f"GLM API 限流。已自动重试 {_LLM_MAX_RETRIES} 次仍未成功。"
            f"当前账号/模型请求过快，请稍后重试，或在 .env 中提高 GLM_RATE_LIMIT_BACKOFF_SECONDS、GLM_MIN_REQUEST_INTERVAL_SECONDS。"
            f"底层错误: {detail}"
        ) from last_error

    raise GLMConnectionError(
        "连接 GLM API 失败。已自动重试 3 次仍未成功。"
        "请检查网络、GLM_API_KEY、模型名，以及是否可访问 https://open.bigmodel.cn 。"
        f"底层错误: {detail}"
    ) from last_error


def _read_text_file(file_path: str) -> str:
    return Path(file_path).read_text(encoding="utf-8").strip()


def _guess_job_input_type(job_path: str) -> str:
    suffix = Path(job_path).suffix.lower()
    if suffix in {".txt", ".md"}:
        return "text"
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp"}:
        return "image"
    raise ValueError("Unsupported job file format. Use .txt/.md or image formats.")


def _image_to_base64(image_path: str) -> str:
    return base64.b64encode(Path(image_path).read_bytes()).decode("utf-8")


def _get_job_requirement_text(job_path: str | None, job_text: str | None, vision_llm: ChatOpenAI) -> str:
    if job_text:
        return job_text.strip()
    if not job_path:
        raise ValueError("You must provide either job_path or job_text")

    if _guess_job_input_type(job_path) == "text":
        return _read_text_file(job_path)

    ext = Path(job_path).suffix.lower().lstrip(".")
    b64 = _image_to_base64(job_path)
    msg = [
        {
            "type": "text",
            "text": (
                "Extract all job requirements from this image in Chinese. "
                "Preserve role summary, responsibilities, must-have skills, preferred skills, and keywords."
            ),
        },
        {"type": "image_url", "image_url": {"url": f"data:image/{ext};base64,{b64}"}},
    ]
    return _invoke_text(vision_llm, [("human", msg)], cache_namespace="vision")


def _convert_doc_to_docx(doc_path: str) -> str:
    src = Path(doc_path)
    out_path = src.with_name(f"{src.stem}_converted.docx")

    # Strategy 1: Microsoft Word COM automation (Windows).
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(src.resolve()))
        document.SaveAs(str(out_path.resolve()), FileFormat=16)
        document.Close(False)
        word.Quit()
        if out_path.exists():
            return str(out_path)
    except Exception:
        pass

    # Strategy 2: LibreOffice soffice headless conversion.
    soffice = shutil.which("soffice")
    if soffice:
        cmd = [
            soffice,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(src.parent.resolve()),
            str(src.resolve()),
        ]
        try:
            subprocess.run(cmd, check=True, capture_output=True, text=True)
            alt_out = src.with_suffix(".docx")
            if alt_out.exists():
                return str(alt_out)
            if out_path.exists():
                return str(out_path)
        except Exception:
            pass

    raise ValueError(
        "上传的是 .doc 文件，但自动转换失败。请安装 Microsoft Word（并可用 pywin32）或 LibreOffice，"
        "或先手动另存为 .docx 后再上传。"
    )


def _ensure_resume_docx(resume_path: str) -> tuple[str, bool]:
    suffix = Path(resume_path).suffix.lower()
    if suffix == ".docx":
        return resume_path, False
    if suffix == ".doc":
        return _convert_doc_to_docx(resume_path), True
    raise ValueError("Resume must be a .docx or .doc file")


def _iter_block_items(parent: _Document | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        parent_elm = getattr(parent, "_tc", None)
        if parent_elm is None:
            parent_elm = getattr(parent, "_element", None)
    if parent_elm is None:
        return

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def _iter_textbox_paragraphs_from_element(element, parent_ctx) -> Iterator[Paragraph]:
    """Yield Paragraph objects from text boxes under the specified xml element."""
    for txbx_content in element.iter(f"{{{_NS_W}}}txbxContent"):
        for p_elm in txbx_content.iterchildren(f"{{{_NS_W}}}p"):
            yield Paragraph(p_elm, parent_ctx)


def _iter_paragraphs_from_container(container, path_prefix: str) -> Iterator[tuple[str, Paragraph]]:
    for block_idx, block in enumerate(_iter_block_items(container), start=1):
        if isinstance(block, Paragraph):
            para_path = f"{path_prefix}/p[{block_idx}]"
            yield para_path, block
            for tb_idx, tb_para in enumerate(_iter_textbox_paragraphs_from_element(block._p, block.part), start=1):
                yield f"{para_path}/textbox[{tb_idx}]", tb_para
            continue

        if isinstance(block, Table):
            for r_idx, row in enumerate(block.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_prefix = f"{path_prefix}/tbl[{block_idx}]/r[{r_idx}]/c[{c_idx}]"
                    yield from _iter_paragraphs_from_container(cell, cell_prefix)


def _iter_document_paragraphs(document: _Document) -> Iterator[tuple[str, Paragraph]]:
    yield from _iter_paragraphs_from_container(document, "body")

    if not _INCLUDE_HEADER_FOOTER:
        return

    seen_header_ids: set[int] = set()
    seen_footer_ids: set[int] = set()
    header_idx = 0
    footer_idx = 0

    for section in document.sections:
        header = section.header
        if id(header._element) not in seen_header_ids:
            seen_header_ids.add(id(header._element))
            header_idx += 1
            yield from _iter_paragraphs_from_container(header, f"header[{header_idx}]")

        footer = section.footer
        if id(footer._element) not in seen_footer_ids:
            seen_footer_ids.add(id(footer._element))
            footer_idx += 1
            yield from _iter_paragraphs_from_container(footer, f"footer[{footer_idx}]")


def _paragraph_hyperlinks(paragraph: Paragraph) -> list[dict[str, str]]:
    links: list[dict[str, str]] = []
    rels = paragraph.part.rels
    for hl in paragraph._p.iter():
        if not str(hl.tag).endswith("}hyperlink"):
            continue
        rid = hl.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", "")
        text_parts = []
        for node in hl.iter():
            if str(node.tag).endswith("}t") and node.text:
                text_parts.append(node.text)
        text = "".join(text_parts)
        target = ""
        if rid and rid in rels:
            target = getattr(rels[rid], "target_ref", "") or ""
        if text or target:
            links.append({"text": text, "target": target})
    return links


def _collect_paragraph_unit(paragraph: Paragraph, unit_id: str, path: str) -> dict[str, Any] | None:
    text = _normalize_text(paragraph.text)
    if not text:
        return None

    runs_meta = []
    for idx, run in enumerate(paragraph.runs):
        runs_meta.append(
            {
                "index": idx,
                "text_len": len(run.text),
                "bold": bool(run.bold) if run.bold is not None else None,
                "italic": bool(run.italic) if run.italic is not None else None,
                "underline": bool(run.underline) if run.underline is not None else None,
                "font_name": run.font.name,
                "font_size": float(run.font.size.pt) if run.font.size is not None else None,
            }
        )

    return {
        "id": unit_id,
        "type": "text",
        "path": path,
        "text": text,
        "style": paragraph.style.name if paragraph.style is not None else "",
        "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
        "runs": runs_meta,
        "hyperlinks": _paragraph_hyperlinks(paragraph),
        "char_limit": max(len(text) + 12, int(len(text) * 1.2)),
    }


def _flatten_nested_text(value: Any) -> list[str]:
    items: list[str] = []
    if isinstance(value, str):
        text = _normalize_text(value)
        if text:
            items.append(text)
        return items
    if isinstance(value, tuple):
        for part in value:
            items.extend(_flatten_nested_text(part))
        return items
    if isinstance(value, list):
        for part in value:
            items.extend(_flatten_nested_text(part))
        return items
    return items


def _extract_supplemental_docx_context(docx_path: str) -> dict[str, Any]:
    if docx2python is None:
        return {"sources": [], "context_text": "", "available": False}

    sources: list[dict[str, Any]] = []
    chunks: list[str] = []
    with docx2python(docx_path) as parsed:
        for name in ("footnotes", "endnotes"):
            lines = _flatten_nested_text(getattr(parsed, name, []))
            if lines:
                sources.append({"name": name, "count": len(lines)})
                chunks.append(f"[{name}]\n" + "\n".join(lines))

        comments = getattr(parsed, "comments", []) or []
        comment_lines: list[str] = []
        for comment in comments:
            parts = [piece for piece in _flatten_nested_text(comment) if piece]
            if parts:
                comment_lines.append(" | ".join(parts))
        if comment_lines:
            sources.append({"name": "comments", "count": len(comment_lines)})
            chunks.append("[comments]\n" + "\n".join(comment_lines))

    return {
        "sources": sources,
        "context_text": "\n\n".join(chunks).strip(),
        "available": True,
    }


def extract_word_layout(docx_path: str) -> dict[str, Any]:
    document = Document(docx_path)
    units: list[dict[str, Any]] = []
    text_units: list[dict[str, Any]] = []
    image_units: list[dict[str, Any]] = []
    full_text_parts: list[str] = []

    text_counter = 1
    for path, para in _iter_document_paragraphs(document):
        unit = _collect_paragraph_unit(para, f"t{text_counter}", path)
        if unit:
            units.append(unit)
            text_units.append(unit)
            full_text_parts.append(unit["text"])
            text_counter += 1

    for i, shape in enumerate(document.inline_shapes, start=1):
        image_units.append(
            {
                "id": f"img{i}",
                "type": "image",
                "width_emu": int(shape.width),
                "height_emu": int(shape.height),
            }
        )

    units.extend(image_units)
    supplemental = _extract_supplemental_docx_context(docx_path)
    return {
        "units": units,
        "text_units": text_units,
        "image_units": image_units,
        "full_text": "\n\n".join(full_text_parts),
        "supplemental_context": supplemental.get("context_text", ""),
        "supplemental_sources": supplemental.get("sources", []),
        "supplemental_available": supplemental.get("available", False),
    }


def _group_resume_sections(text_units: list[dict[str, Any]]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    current = {"title": "未命名模块", "unit_ids": [], "content": []}

    for unit in text_units:
        text = unit["text"].strip()
        style = unit.get("style", "")
        is_heading = ("Heading" in style) or (len(text) <= 20 and any(k in text for k in ["教育", "项目", "经历", "技能", "证书", "评价"]))
        if is_heading:
            if current["unit_ids"]:
                sections.append({**current, "content": "\n".join(current["content"]).strip()})
            current = {"title": text, "unit_ids": [unit["id"]], "content": [unit["text"]]}
            continue

        current["unit_ids"].append(unit["id"])
        current["content"].append(unit["text"])

    if current["unit_ids"]:
        sections.append({**current, "content": "\n".join(current["content"]).strip()})

    return sections


def _analyze_job_requirement(text_llm: ChatOpenAI, job_requirement: str) -> dict[str, Any]:
    prompt = (
        "你是招聘需求分析助手。请将岗位要求结构化。\n"
        "输出 JSON：{\"target_role\":\"\",\"core_requirements\":[],\"keywords\":[],\"priorities\":[],\"nice_to_have\":[]}。\n"
        "不要输出解释。\n\n"
        f"岗位原文：\n{job_requirement}"
    )
    return _invoke_json(text_llm, prompt)


def _build_matching_evidence(text_llm: ChatOpenAI, structured_job: dict[str, Any], sections: list[dict[str, Any]]) -> dict[str, Any]:
    prompt = (
        "你是简历匹配分析助手。请输出岗位要求与简历模块的对应关系。\n"
        "输出 JSON：{\"matches\":[{\"requirement\":\"\",\"section_titles\":[],\"reason\":\"\",\"missing\":false}]}。\n"
        "不要输出解释。\n\n"
        f"岗位结构化：\n{json.dumps(structured_job, ensure_ascii=False)}\n\n"
        f"简历模块：\n{json.dumps(sections, ensure_ascii=False)}"
    )
    return _invoke_json(text_llm, prompt)


def _apply_char_limit(text: str, char_limit: int) -> str:
    t = _normalize_text(text)
    if len(t) <= char_limit:
        return t
    return t[: max(0, char_limit - 1)].rstrip() + "…"


def _optimize_text_units(
    text_llm: ChatOpenAI,
    text_units: list[dict[str, Any]],
    job_requirement: str,
    full_resume_text: str = "",
    structured_job: dict[str, Any] | None = None,
    evidence_map: dict[str, Any] | None = None,
    revision_advice: list[str] | None = None,
    supplemental_context: str = "",
) -> list[dict[str, Any]]:
    payload = [
        {
            "id": u["id"],
            "style": u.get("style", ""),
            "path": u.get("path", ""),
            "char_limit": u["char_limit"],
            "text": u["text"],
        }
        for u in text_units
    ]
    revise = "\n".join(f"- {x}" for x in (revision_advice or [])) or "无"
    structured_job_text = json.dumps(structured_job or {}, ensure_ascii=False)
    evidence_map_text = json.dumps(evidence_map or {}, ensure_ascii=False)
    context_block = f"补充上下文（脚注/尾注/批注）：\n{supplemental_context}\n\n" if supplemental_context else ""
    full_resume_block = f"完整简历：\n{full_resume_text}\n\n" if full_resume_text else ""
    prompt = (
        "你是资深中文简历优化专家。\n"
        "在不改变真实信息的前提下，按岗位要求优化简历。\n"
        "必须保持模块顺序不变；每个 id 都要返回；不要输出解释。\n"
        "输出 JSON：{\"blocks\":[{\"id\":\"\",\"optimized_text\":\"\"}]}。\n\n"
        f"岗位需求：\n{job_requirement}\n\n"
        f"岗位结构化：\n{structured_job_text}\n\n"
        f"岗位-证据映射：\n{evidence_map_text}\n\n"
        f"修订意见：\n{revise}\n\n"
        f"{context_block}"
        f"{full_resume_block}"
        f"待优化单元：\n{json.dumps(payload, ensure_ascii=False)}"
    )

    data = _invoke_json(text_llm, prompt)
    optimized_map = {item["id"]: _normalize_text(item.get("optimized_text", "")) for item in data.get("blocks", []) if item.get("id")}

    missing = [u["id"] for u in text_units if u["id"] not in optimized_map]
    if missing:
        raise ValueError(f"Model output missing unit ids: {missing[:5]}")

    out = []
    for unit in text_units:
        optimized_text = _apply_char_limit(optimized_map.get(unit["id"], unit["text"]), int(unit["char_limit"]))
        out.append({**unit, "optimized_text": optimized_text})
    return out


def _verify_optimization(
    text_llm: ChatOpenAI,
    job_requirement: str,
    original_units: list[dict[str, Any]],
    optimized_units: list[dict[str, Any]],
) -> dict[str, Any]:
    original_text = "\n\n".join(u["text"] for u in original_units)
    optimized_text = "\n\n".join(u["optimized_text"] for u in optimized_units)
    prompt = (
        "你是简历优化质检器。检查是否体现岗位匹配、成果表达提升、且无编造。\n"
        "输出 JSON：{\"passed\":true,\"issues\":[],\"revision_advice\":[]}。\n"
        "不要输出解释。\n\n"
        f"岗位要求：\n{job_requirement}\n\n"
        f"原文：\n{original_text}\n\n"
        f"优化后：\n{optimized_text}"
    )
    return _invoke_json(text_llm, prompt)


def _iter_target_paragraphs(document: Document) -> list[Paragraph]:
    return [para for _, para in _iter_document_paragraphs(document)]


def _rewrite_paragraph_preserve_runs(paragraph: Paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return

    lengths = [len(run.text) for run in runs]
    total = sum(lengths)
    if total == 0:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
        return

    idx = 0
    for run, size in zip(runs, lengths):
        if idx >= len(new_text):
            run.text = ""
            continue
        chunk = new_text[idx : idx + size]
        run.text = chunk
        idx += len(chunk)

    if idx < len(new_text):
        runs[-1].text = runs[-1].text + new_text[idx:]


def save_optimized_resume_docx(
    optimized_units: list[dict[str, Any]],
    source_resume_path: str,
    output_docx_path: str,
) -> None:
    document = Document(source_resume_path)
    paragraphs = _iter_target_paragraphs(document)

    unit_map = {unit["id"]: unit for unit in optimized_units}
    seq = 0
    for paragraph in paragraphs:
        raw = _normalize_text(paragraph.text)
        if not raw:
            continue
        seq += 1
        uid = f"t{seq}"
        if uid not in unit_map:
            continue
        optimized_text = unit_map[uid]["optimized_text"]
        optimized_text = _preserve_leading_indent(paragraph.text, optimized_text)
        _rewrite_paragraph_preserve_runs(paragraph, optimized_text)

    out_path = Path(output_docx_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))


def save_word_units_snapshot(word_units: list[dict[str, Any]], output_json_path: str) -> None:
    out_path = Path(output_json_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(word_units, ensure_ascii=False, indent=2), encoding="utf-8")


def optimize_resume_docx(
    resume_path: str,
    job_path: str | None = None,
    job_text: str | None = None,
) -> dict[str, Any]:
    load_dotenv()

    if not os.environ.get("GLM_API_KEY"):
        raise EnvironmentError("GLM_API_KEY is not set. Please configure it in .env")
    if not Path(resume_path).exists():
        raise FileNotFoundError(f"Resume file not found: {resume_path}")
    if not job_path and not job_text:
        raise ValueError("You must provide either job_path or job_text")
    if job_path and job_text:
        raise ValueError("Please provide only one: job_path or job_text")

    normalized_resume_path, converted_from_doc = _ensure_resume_docx(resume_path)

    text_model = os.environ.get("GLM_TEXT_MODEL", "glm-4-plus")
    vision_model = os.environ.get("GLM_VISION_MODEL", "glm-4v-plus-0111")
    text_llm = _build_llm(text_model)
    vision_llm = _build_llm(vision_model, temperature=0)

    layout = extract_word_layout(normalized_resume_path)
    if not layout["text_units"]:
        raise ValueError("未能从 Word 简历中提取到有效文字内容")

    job_requirement = _get_job_requirement_text(job_path, job_text, vision_llm)
    sections = _group_resume_sections(layout["text_units"])
    structured_job: dict[str, Any] = {}
    evidence_map: dict[str, Any] = {}

    if _WORKFLOW_MODE == "full":
        structured_job = _analyze_job_requirement(text_llm, job_requirement)
        evidence_map = _build_matching_evidence(text_llm, structured_job, sections)
    else:
        structured_job = {
            "workflow_mode": "lite",
            "section_titles": [section["title"] for section in sections],
        }
        evidence_map = {
            "workflow_mode": "lite",
            "note": "Skipped extra GLM analysis calls to reduce rate-limit risk.",
        }

    optimized_units = _optimize_text_units(
        text_llm,
        layout["text_units"],
        job_requirement,
        full_resume_text=layout["full_text"] if _WORKFLOW_MODE == "full" else "",
        structured_job=structured_job,
        evidence_map=evidence_map,
        supplemental_context=layout.get("supplemental_context", ""),
    )

    verification: dict[str, Any] = {
        "passed": True,
        "skipped": not _ENABLE_VERIFICATION,
        "reason": "GLM_ENABLE_VERIFICATION=false，为降低 GLM 限流风险已跳过质检。" if not _ENABLE_VERIFICATION else "",
    }
    if _ENABLE_VERIFICATION:
        try:
            verification = _verify_optimization(text_llm, job_requirement, layout["text_units"], optimized_units)
            if not verification.get("passed", True):
                optimized_units = _optimize_text_units(
                    text_llm,
                    layout["text_units"],
                    job_requirement,
                    full_resume_text=layout["full_text"] if _WORKFLOW_MODE == "full" else "",
                    structured_job=structured_job,
                    evidence_map=evidence_map,
                    revision_advice=verification.get("revision_advice", []) or verification.get("issues", []),
                    supplemental_context=layout.get("supplemental_context", ""),
                )
        except GLMRateLimitError as exc:
            verification = {
                "passed": True,
                "skipped": True,
                "reason": f"质检阶段触发 GLM 限流，已保留当前优化结果。{exc}",
            }

    return {
        "word_units": layout["units"],
        "resume_text": layout["full_text"],
        "job_requirement": job_requirement,
        "structured_job": structured_job,
        "evidence_map": evidence_map,
        "optimized_units": optimized_units,
        "verification": verification,
        "resolved_resume_path": normalized_resume_path,
        "converted_from_doc": converted_from_doc,
        "parse_report": {
            "text_unit_count": len(layout["text_units"]),
            "image_unit_count": len(layout["image_units"]),
            "supplemental_available": layout.get("supplemental_available", False),
            "supplemental_sources": layout.get("supplemental_sources", []),
            "workflow_mode": _WORKFLOW_MODE,
            "cache_enabled": _ENABLE_GLM_CACHE,
        },
    }
